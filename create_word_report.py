from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

doc = Document()

# Titre
title = doc.add_heading('Moteur de Recommandation - Micro-Aventures', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Auteurs
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Auteurs : ').bold = True
p.add_run('Zahra Kodia Aouina & Siwar Gorrab\n')
p.add_run('Classe : ').bold = True
p.add_run('2 LNBI - ISG TUNIS')

# Section 1
add_heading(doc, '1. Problématique', 1)
doc.add_paragraph("Avec la multiplication des offres de loisirs et de tourisme local, les utilisateurs se retrouvent souvent perdus face à une surabondance d'options (surcharge d'information). L'objectif de ce projet est de concevoir et développer un moteur de recommandation intelligent capable de suggérer des \"Micro-Aventures et Expériences Insolites\" adaptées aux préférences uniques de chaque utilisateur.")
doc.add_paragraph("Le défi technique réside dans la nécessité de combiner plusieurs approches de filtrage pour pallier les défauts individuels de chaque méthode et d'offrir une interface explicable et transparente.")

# Section 2
add_heading(doc, '2. Méthodologie', 1)
doc.add_paragraph("Notre système repose sur une base de données PostgreSQL personnalisée (sr) contenant trois tables principales (Produit, Users, Notes). Nous avons développé un modèle hybride combinant trois algorithmes :")

add_heading(doc, 'A. Filtrage basé sur le contenu (Content-Based - TD4)', 2)
doc.add_paragraph("Nous analysons le texte des descriptions des expériences pour recommander des items similaires à ceux que l'utilisateur a déjà bien notés.", style='List Bullet')
doc.add_paragraph("NLP : Utilisation de NLTK pour la tokenisation, le stemming (FrenchStemmer) et la suppression des mots vides.", style='List Bullet')
doc.add_paragraph("Représentation : Création d'une Matrice Binaire associant chaque produit à ses racines de mots.", style='List Bullet')
doc.add_paragraph("Similarité : Calcul de la similarité Cosinus entre les vecteurs des produits.", style='List Bullet')

add_heading(doc, 'B. Filtrage Collaboratif (User-Based - TD5)', 2)
doc.add_paragraph("Nous exploitons la sagesse de la foule en recommandant ce qu'ont aimé des utilisateurs similaires.", style='List Bullet')
doc.add_paragraph("Matrice Utilisateur-Item : Création d'un pivot à partir de l'historique des notes.", style='List Bullet')
doc.add_paragraph("Similarité : Calcul de la similarité Cosinus entre les utilisateurs.", style='List Bullet')
doc.add_paragraph("Prédiction : Pour un produit non noté, calcul d'une note estimée basée sur la moyenne pondérée des notes des \"K Plus Proches Voisins\" (KNN).", style='List Bullet')

add_heading(doc, 'C. Recommandation Temporelle (Time-Aware)', 2)
doc.add_paragraph("Décroissance Temporelle : Les notes récentes ont un poids mathématique plus fort (fonction exponentielle) que les notes anciennes.", style='List Bullet')
doc.add_paragraph("Saisonnalité : Les activités correspondant à la saison actuelle reçoivent un boost de pertinence.", style='List Bullet')

add_heading(doc, 'D. Modèle Hybride et Explicabilité', 2)
doc.add_paragraph("Les scores des trois algorithmes sont combinés (moyenne pondérée). Le système détermine quel algorithme a le plus contribué au score final pour afficher une explication claire à l'utilisateur.")

# Section 3
add_heading(doc, '3. Résultats et Évaluation (Bonus)', 1)
doc.add_paragraph("L'application développée en Streamlit offre une interface ergonomique et dynamique. Pour évaluer la pertinence, nous avons implémenté les métriques classiques :")
doc.add_paragraph("Precision@K : Mesure la proportion d'expériences pertinentes parmi le Top-K recommandé.", style='List Bullet')
doc.add_paragraph("Recall@K : Mesure la capacité du système à retrouver les bonnes expériences.", style='List Bullet')
doc.add_paragraph("NDCG@K : Qualité du classement.", style='List Bullet')
doc.add_paragraph("Nos tests sur le dataset synthétique montrent une excellente capacité du modèle hybride à s'adapter aux profils variés, justifiant pleinement l'intégration des 3 approches.")

doc.save('Livrable_3_Rapport_Projet.docx')
print("Document Word généré avec succès !")
